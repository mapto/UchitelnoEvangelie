#!/usr/bin/env python3

from typing import List, Dict
import logging as log

from pathlib import Path

from docx import Document  # type: ignore
from docx.table import _Cell  # type: ignore
from lxml import etree  # type: ignore

from const import LINE_CH
from schema import ns
from schema import tag_paragraph, tag_run, tag_text, tag_br
from schema import tag_commentStart, tag_commentEnd
from model import Index, Word, Comment, WordList
from util import Buffer


def parse_comments(doc: Document) -> Dict[int, Comment]:
    """Parse the comments body"""
    xml = None
    result: Dict[int, Comment] = {}
    for part in doc.part.package.parts:
        if part.partname == "/word/comments.xml":
            xml = part.blob
    if not xml:
        return result

    root = etree.fromstring(xml)
    for next in root.getchildren():
        c = Comment.fromXml(next)
        result[c.id] = c

    return result


def parse_page(
    ch: int, page: str, rows: List[int], cell: _Cell, comments: Dict[int, Comment]
) -> WordList:
    """Parse the actual running text, annotated with comment references"""
    buffer = Buffer()
    result = WordList()
    row = rows.pop(0)
    for i, par in enumerate(cell._element):
        if par.tag != tag_paragraph:
            continue
        for sec in par:
            if sec.tag == tag_run:
                for content in sec:
                    if content.tag == tag_text:
                        buffer.add(content.text)
                    elif content.tag == tag_br:
                        buffer.swap()
                        idx = Index(ch, page, row)
                        compiled = buffer.compile_buffer(idx, comments)
                        buffer.build_context(compiled)
                        result += buffer.flush()
                        row = rows.pop(0)
                        if not rows:  # add row numbers that might not be provided
                            rows.append(row + 1)

            elif sec.tag == tag_commentStart:
                idx = Index(ch, page, row)
                buffer.line_words += buffer.compile_buffer(idx, comments)
                buffer.swap_clean()

                id = int(sec.xpath("./@w:id", namespaces=ns)[0])
                buffer.comments.add(id)

            elif sec.tag == tag_commentEnd:
                id = int(sec.xpath("./@w:id", namespaces=ns)[0])
                idx = Index(ch, page, row)
                compiled = buffer.compile_buffer(idx, comments)
                compiled[-1].variant = compiled[-1].variant.replace(
                    LINE_CH, comments[id].annotation
                )

                buffer.line_words += compiled
                buffer.swap_clean()

                # additions
                for c in buffer.comments:
                    for comment in comments[c].addition:
                        idx = Index(ch, page, row)
                        buffer.line_words += Word(idx, variant=f"+{comment}")

                buffer.comments.remove(id)

        buffer.swap()
        idx = Index(ch, page, row)
        log.debug(idx)
        if i < len(cell._element) - 1:  # if not end of page, behave like line break
            compiled = buffer.compile_buffer(idx, comments)
        else:
            compiled = buffer.compile_words(idx, buffer.extract_comment(comments))
        buffer.build_context(compiled)
        result += buffer.flush()
        if not rows:
            break
        row = rows.pop(0)
        if not rows:  # add rows that might not be provided
            rows.append(row + 1)

    # TODO: Expand comment selection to cover whole words

    # at end all comment starts should be matched by comment ends within the page
    assert not buffer.comments

    return result


def parse_document(ch: int, doc: Document, comments: Dict[int, Comment]) -> WordList:
    """Parses a particular type of documents with content in a 2x2 table"""
    words = WordList()
    for page in doc.tables:
        page_name = page.cell(0, 0).text
        page_rows_nums = [int(r) for r in page.cell(1, 0).text.split("\n") if r]
        cell = page.cell(1, 1)
        page_words = parse_page(ch, page_name, page_rows_nums, cell, comments)
        words += page_words
    return words


def import_chapter(fname: str) -> WordList:
    """Import a preformatted chapter of a manuscript"""
    doc = Document(fname)
    parts = Path(fname).name.split("-")
    book_prefix = 0
    try:
        book_prefix = int(parts[0])
    except ValueError:
        log.error("Името на файла не започва с <номер>-, затова допускаме глава 0.")

    comments = parse_comments(doc)
    log.debug(comments)
    return parse_document(book_prefix, doc, comments)


if __name__ == "__main__":
    # import_lines("../text/01-slovo1-tab.docx")
    # book_lines = import_lines("../text/00-Prolog-tab.docx")
    # print(import_lines("../new/01-slovo1-tab.docx"))
    # print(import_comments(Document('../new/01-slovo1-tab.docx')))
    # print(locate_comments(Document('../new/01-slovo1-tab.docx').tables[0].cell(1,1)._element[1]))
    import_chapter("test/01-slovo1-4b.docx")
