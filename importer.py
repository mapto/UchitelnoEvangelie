#!/usr/bin/env python3

from typing import List, Tuple, Dict, Set

import re

from docx import Document  # type: ignore
from docx.table import _Cell  # type: ignore
from lxml import etree, html  # type: ignore

from schema import ns
from schema import (
    tag_paragraph,
    tag_run,
    tag_commentStart,
    tag_commentEnd,
    tag_text,
    tag_br,
)
from model import Index, Word, Comment


def parse_comments(doc: Document) -> Dict[int, Comment]:
    """Parse the comments body"""
    xml = None
    result: Dict[int, Comment] = {}
    for part in doc.part.package.parts:
        if part.partname == "/word/comments.xml":
            xml = part.blob
    if not xml:
        return result

    root = etree.fromstring(xml)
    for next in root.getchildren():
        c = Comment.fromXml(next)
        result[c.id] = c

    return result


def _compile_words(idx: Index, buffer: str, comment: str) -> List[Word]:
    """Convert a text into a list of annotated words"""
    result: List[Word] = []
    words = re.split(r"\s", buffer)
    for w in words:
        new_word = Word(idx, w, buffer, variant=comment)
        if result:
            new_word.appendTo(result[-1])
        result.append(new_word)
    return result


def _compile_buffer(
    idx: Index, buffer: str, comments: Dict[int, Comment], current: Set[int]
):
    parts = set()
    for c in current:
        if comments[c].annotation:
            if comments[c].annotation.startswith("om"):
                parts.add(comments[c].annotation)
            else:
                parts.add("↓")
    comment = ",".join(parts)
    return _compile_words(idx, buffer, comment)


def _merge(head: List[Word], tail: List[Word]) -> List[Word]:
    """Returns first parameter"""
    if head:
        head[-1].prependTo(tail[0])
    head.extend(tail)
    return head


def parse_page(
    ch: int, page: str, rows: List[int], cell: _Cell, comments: Dict[int, Comment]
) -> List[Word]:
    """Parse the actual running text, annotated with comment references"""
    result: List[Word] = []
    open_comments: Set[int] = set()
    buff = ""
    line = ""
    line_words: List[Word] = []
    row = rows.pop(0)
    for par in cell._element:
        if par.tag != tag_paragraph:
            continue
        for sec in par:
            if sec.tag == tag_run:
                for content in sec:
                    if content.tag == tag_text:
                        buff += content.text
                    elif content.tag == tag_br:
                        line += buff
                        idx = Index(ch, page, row)
                        compiled = _compile_buffer(idx, buff, comments, open_comments)
                        line_words = _merge(line_words, compiled)
                        for w in line_words:
                            w.line_context = line
                        result = _merge(result, line_words)
                        buff = ""
                        line = ""
                        line_words = []
                        row = rows.pop(0)
                        if not rows:
                            rows.append(row + 1)

            elif sec.tag == tag_commentStart:
                idx = Index(ch, page, row)
                compiled = _compile_buffer(idx, buff, comments, open_comments)
                line_words = _merge(line_words, compiled)
                line += buff
                buff = ""

                id = int(sec.xpath("./@w:id", namespaces=ns)[0])
                open_comments.add(id)

            elif sec.tag == tag_commentEnd:
                id = int(sec.xpath("./@w:id", namespaces=ns)[0])
                idx = Index(ch, page, row)
                compiled = _compile_buffer(idx, buff, comments, open_comments)
                compiled[-1].variant = compiled[-1].variant.replace(
                    "↓", comments[id].annotation
                )
                line_words = _merge(line_words, compiled)
                line += buff
                buff = ""

                # additions
                for c in open_comments:
                    for comment in comments[c].addition:
                        idx = Index(ch, page, row)
                        new_word = Word(idx, variant=f"+{comment}")
                        if line_words:
                            new_word.appendTo(line_words[-1])
                        line_words.append(new_word)

                open_comments.remove(id)

    line += buff
    comment_parts = set([comments[c].annotation for c in open_comments])
    comment = ",".join(comment_parts)
    idx = Index(ch, page, row)
    compiled = _compile_words(idx, buff, comment)
    line_words = _merge(line_words, compiled)
    for w in line_words:
        w.line_context = line
    result = _merge(result, line_words)

    # TODO: Expand comment selection to cover whole words
    row = rows.pop(0)
    if not rows:
        rows.append(row + 1)

    # at end all comment starts should be matched by comment ends
    assert not open_comments

    return result


def parse_document(ch: int, doc: Document, comments: Dict[int, Comment]) -> List[Word]:
    """Parses a particular type of documents with content in a 2x2 table"""
    words: List[Word] = []
    for page in doc.tables:
        page_name = page.cell(0, 0).text
        page_rows_nums = [int(r) for r in page.cell(1, 0).text.split("\n") if r]
        cell = page.cell(1, 1)
        page_words = parse_page(ch, page_name, page_rows_nums, cell, comments)
        words = _merge(words, page_words)
    return words


def import_chapter(fname: str) -> List[Word]:
    """Import a preformatted chapter of a manuscript"""
    doc = Document(fname)
    book_prefix = int(fname.split("/")[-1].split("-")[0])
    comments = parse_comments(doc)
    # print(comments)
    return parse_document(book_prefix, doc, comments)


if __name__ == "__main__":
    # import_lines("../text/01-slovo1-tab.docx")
    # book_lines = import_lines("../text/00-Prolog-tab.docx")
    # print(import_lines("../new/01-slovo1-tab.docx"))
    # print(import_comments(Document('../new/01-slovo1-tab.docx')))
    # print(locate_comments(Document('../new/01-slovo1-tab.docx').tables[0].cell(1,1)._element[1]))
    import_chapter("test/01-slovo1-4b.docx")
