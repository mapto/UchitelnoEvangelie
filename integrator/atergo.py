"""The exporter specific to the indexgenerator"""
from typing import Union
import logging as log

from sortedcontainers import SortedDict, SortedSet  # type: ignore

from docx import Document  # type: ignore
from docx.shared import Cm  # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore

from util import ord_word

from wordproc import _generate_text
from wordproc import GENERIC_FONT, fonts
from model import Counter


def _generate_counts(par, d: Union[SortedDict, dict], trans: bool = False) -> None:
    try:
        c = Counter.get_dict_counts(d).get_counts(trans)
    except StopIteration:
        keys = []
        key = next(iter(d))
        if key:
            keys += [key]
        while key and key in d and bool(d[key]):
            d = d[key]
            key = next(iter(d))
            if key:
                keys += [key]
        # TODO: Better reporting
        if keys:
            log.error(f"При генериране възникна проблем с {'|'.join(keys)}")
        else:
            log.error(f"При генериране възникна неидентифициран проблем")
        log.error(f"При генериране неуспешно преброяване в {d}")
        return
    run = par.add_run()
    run.add_text("(")
    if c[0]:
        run.add_text(str(c[0]))
        if c[1]:
            run.add_text(" + ")
    if c[1]:
        run.add_text(str(c[1]))
        _generate_text(par, "var", superscript=True)
    _generate_text(par, ") ")


def _generate_line(lang: str, d: SortedDict, doc: Document) -> None:
    """Builds the (non-hierarchical) entries of the atergo dictionary.

    Args:
        lang (str): original language
        d (SortedDict): level of dictionary to be exported
        doc (Document): export document
    """
    for rli, (li, next_d) in d.items():
        # if not li:
        #     continue
        if li:
            par = doc.add_paragraph()
            par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # pylint: disable=E1101
            par.style.font.name = GENERIC_FONT
            par.paragraph_format.space_before = Cm(0)
            par.paragraph_format.space_after = Cm(0)

            _generate_counts(par, next_d, False)
            _generate_text(par, li, fonts[lang])


def generate_docx(d: SortedDict, lang: str, fname: str) -> None:
    atergo_d = SortedDict(ord_word)
    for k, v in d.items():
        parts = k.split(" ")
        if len(parts) < 2:
            lemma = k
        else:
            lemma = f'{" ".join(parts[1:])} {parts[0]}'
        atergo_d[lemma[::-1]] = (lemma, v)

    doc = Document()
    doc.styles["Normal"].font.name = GENERIC_FONT
    _generate_line(lang, atergo_d, doc)
    doc.save(fname)
