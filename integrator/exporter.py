"""The exporter specific to the integrator"""

from typing import Dict, Tuple
import logging as log

from sortedcontainers import SortedDict, SortedSet  # type: ignore

from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore

from const import INDENT_CH, SPECIAL_CHARS
from config import FROM_LANG, TO_LANG, other_lang
from const import CF_SEP
from const import BRACE_OPEN, BRACE_CLOSE

from util import main_source, subscript
from model import Source, Alignment, Alternative, Usage

from wordproc import _generate_text, any_grandchild
from wordproc import GENERIC_FONT, fonts, colors


def _generate_usage_alt_main(par, uc: Usage, source_name: str) -> None:
    alt = uc.main_alt
    _generate_text(par, " ")
    prefix = f"{alt.semantic} " if alt.semantic else ""
    _generate_text(par, f"{prefix}{alt.word}", fonts[uc.lang])
    _generate_text(par, f" {source_name}")
    if alt.cnt > 1:
        _generate_text(par, subscript(alt.cnt, uc.lang), subscript=True)


def _generate_usage_alt_vars(
    par, u: Alignment, lang: str, alt_var: Dict[Source, Alternative]
) -> None:
    # Word is stored in Index only for orig.
    # Thus the check u.lang == lang
    if u.orig.lang == lang and not any(
        v.word not in u.orig.word for v in alt_var.values()
    ):
        return

    first = True
    _generate_text(par, f" {BRACE_OPEN[lang]}")
    for alt in alt_var.values():
        word, cnt = alt.word, alt.cnt
        if u.orig.lang == lang and word in u.orig.word:
            continue
        if first:
            first = False
        else:
            _generate_text(par, ", ")
        prefix = f"{alt.semantic} " if alt.semantic else ""
        _generate_text(par, f"{prefix}{word}", fonts[lang])
        subs = subscript(cnt, lang)
        if subs:
            _generate_text(par, subs, subscript=True)
    _generate_text(par, BRACE_CLOSE[lang])


def _generate_index(par, u: Alignment) -> None:
    s = str(u.idx)
    _generate_text(par, s, bold=u.bold, italic=u.italic)

    sl_cnt = u.orig.cnt if u.orig.lang == FROM_LANG else u.trans.cnt
    gr_cnt = u.trans.cnt if u.orig.lang == FROM_LANG else u.orig.cnt

    other_before_own = False
    if u.var() and u.var().has_lang(FROM_LANG):
        if not u.var().has_lang(TO_LANG):
            if gr_cnt > 1:
                _generate_text(par, subscript(gr_cnt, TO_LANG), subscript=True)
            other_before_own = True
        _generate_text(par, str(u.var().by_lang(FROM_LANG)), superscript=True)
    if sl_cnt > 1:
        _generate_text(par, subscript(sl_cnt, FROM_LANG), subscript=True)
    elif u.var() and u.var().has_lang(FROM_LANG) and u.var().has_lang(TO_LANG):
        _generate_text(par, "-", superscript=True)
    if not other_before_own:
        if u.var() and u.var().has_lang(TO_LANG):
            _generate_text(par, str(u.var().by_lang(TO_LANG)), superscript=True)
        if gr_cnt > 1:
            _generate_text(par, subscript(gr_cnt, TO_LANG), subscript=True)


def _generate_usage(par, u: Alignment) -> None:
    _generate_index(par, u)
    if not u.has_alternatives():
        return

    _generate_text(par, f" {CF_SEP}")
    if u.orig.main_alt.word:
        source = main_source(u.orig.lang, u.idx.data[2] == "W")
        _generate_usage_alt_main(par, u.orig, source)

    if u.orig.var_alt:
        _generate_usage_alt_vars(par, u, u.orig.lang, u.orig.var_alt)

    # previous addition certainly finished with GENERIC_FONT
    if u.trans.main_alt.word:
        source = main_source(u.trans.lang, u.idx.data[2] == "W")
        _generate_usage_alt_main(par, u.trans, source)

    if u.trans.var_alt:
        _generate_usage_alt_vars(par, u, u.trans.lang, u.trans.var_alt)


def docx_usage(par, key: Tuple[str, str], usage: SortedSet, src_style: str) -> None:
    """
    key: (word,translation)
    usage: list of indices of usages also containing their styles
    """
    other_style = other_lang[src_style]

    _generate_text(par, key[0], fonts[src_style], colors[src_style])
    _generate_text(par, "/")
    _generate_text(par, key[1], fonts[other_style], colors[other_style])
    _generate_text(par, " (")

    first = True
    for next in usage:
        if first:
            first = False
        else:
            _generate_text(par, ", ")
        _generate_usage(par, next)
    _generate_text(par, ")")


def _generate_usage_line(lang: str, d: SortedDict, doc: Document) -> None:
    """Lists ordered occurences for each word usage pair"""
    trans_lang = TO_LANG if lang == FROM_LANG else FROM_LANG
    for t, bottom_d in d.items():
        par = doc.add_paragraph()
        par.style.font.name = GENERIC_FONT
        par.paragraph_format.left_indent = Pt(30)
        par.paragraph_format.first_line_indent = Pt(-10)

        _generate_text(par, t, fonts[trans_lang])

        run = par.add_run()
        run.add_text(": ")
        first = True
        pairs = SortedDict(bottom_d.items())
        for key in pairs:
            usage = bottom_d[key]
            if not first:
                par.add_run().add_text("; ")
            try:
                docx_usage(par, key, usage, lang)
            except Exception as e:
                log.error(
                    f"При експортиране възникна проблем в ред {usage.idx} ({key}) или групата му"
                )
                log.error(e)
                continue
            first = False


def _export_line(level: int, lang: str, d: SortedDict, doc: Document):
    """Builds the hierarchical entries for detailed comparison.
    Recursion ensures that this works with variable depth.

    Args:
        level (int): 0-indexed depth, runs along with dict depth
        lang (str): original language
        d (SortedDict): level of dictionary to be exported
        doc (Document): export document
    """
    for li, next_d in d.items():
        if level == 0 and not li:
            continue
        if li:
            par = doc.add_paragraph()
            par.style.font.name = GENERIC_FONT
            if level > 0:
                par.paragraph_format.first_line_indent = Pt(10)

            prefix = "" if li[0] in SPECIAL_CHARS else f"{INDENT_CH} " * level
            _generate_text(
                par, f"{prefix}{li}", fonts[lang], size=Pt(14 if level == 0 else 12)
            )

        try:
            any_of_any = any_grandchild(next_d)
        except AssertionError as ae:
            suffix = f", част от {d}" if not li else ""
            log.error(
                f"При обособяване на речника възникна проблем с една от {len(next_d)} употреби на '{li}' "
                f"в лема на ниво {level}{suffix}. "
                "Тази лема ще бъде прескочена."
            )
            log.error(ae)
            continue

        if type(any_of_any) is SortedSet:  # bottom of structure
            _generate_usage_line(lang, next_d, doc)
        else:
            try:
                _export_line(level + 1, lang, next_d, doc)
            except AssertionError:
                suffix = f", част от {d}" if not li else ""
                log.error(
                    f"При генериране на ред в речника възникна проблем с една от {len(next_d)} употреби на {li}"
                    f"в лема номер {level}{suffix}. "
                    "Тази лема ще бъде прескочена."
                )


def export_docx(d: SortedDict, lang: str, fname: str) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = GENERIC_FONT
    _export_line(0, lang, d, doc)

    doc.save(fname)
