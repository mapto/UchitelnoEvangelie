"""The exporter specific to the indexgenerator"""
from typing import Union
import logging as log

from sortedcontainers import SortedDict, SortedSet  # type: ignore

from docx import Document  # type: ignore
from docx.shared import Pt, Cm  # type: ignore

from config import FROM_LANG, TO_LANG
from const import INDENT_CH, SPECIAL_CHARS
from const import CF_SEP
from const import BRACE_OPEN, BRACE_CLOSE
from util import main_source, subscript
from model import Alignment, Usage

from wordproc import _generate_text, any_grandchild
from wordproc import GENERIC_FONT, fonts
from model import Counter

BULLET_STYLE = "List Bullet"
LEVEL_OFFSET = 0.4


def _generate_usage_alt_main(par, uc: Usage, source_name: str) -> None:
    alt = uc.main_alt
    _generate_text(par, " ")
    _generate_text(par, alt.lemma(), fonts[uc.lang])
    _generate_text(par, f" {source_name}")
    if alt.cnt > 1:
        _generate_text(par, subscript(alt.cnt, uc.lang), subscript=True)


def _generate_usage_alt_vars(par, uc: Usage) -> None:
    first = True
    _generate_text(par, f" {BRACE_OPEN[uc.lang]}")
    for lsrc, alt in uc.var_alt.items():
        # for lsrc, lemma in uc.alt.var_lemmas.items():
        args = [
            # tpl[1] for wsrc, tpl in alt_var.var_words.items() if wsrc and tpl[0] and wsrc.inside([lsrc])
            tpl.cnt
            for wsrc, tpl in uc.var_alt.items()
            if wsrc.inside([lsrc])
        ]
        cnt = max(args) if args else 1
        if first:
            first = False
        else:
            _generate_text(par, ", ")
        _generate_text(par, alt.lemma(), fonts[uc.lang])
        if lsrc:
            _generate_text(par, str(lsrc._sort_vars()), superscript=True)
        if cnt > 1:
            _generate_text(par, subscript(cnt, uc.lang), subscript=True)
    _generate_text(par, BRACE_CLOSE[uc.lang])


def _generate_index(par, u: Alignment) -> None:
    s = str(u.idx)
    _generate_text(par, s, bold=u.bold, italic=u.italic)

    sl_cnt = u.orig.cnt if u.orig.lang == FROM_LANG else u.trans.cnt
    gr_cnt = u.trans.cnt if u.orig.lang == FROM_LANG else u.orig.cnt

    other_before_own = False
    if u.var() and u.var().has_lang(FROM_LANG):
        if not u.var().has_lang(TO_LANG):
            if gr_cnt > 1:
                _generate_text(par, subscript(gr_cnt, TO_LANG), subscript=True)
            other_before_own = True
        _generate_text(par, str(u.var().by_lang(FROM_LANG)), superscript=True)
    if sl_cnt > 1:
        _generate_text(par, subscript(sl_cnt, FROM_LANG), subscript=True)
    elif u.var() and u.var().has_lang(FROM_LANG) and u.var().has_lang(TO_LANG):
        _generate_text(par, "-", superscript=True)
    if not other_before_own:
        if u.var() and u.var().has_lang(TO_LANG):
            _generate_text(par, str(u.var().by_lang(TO_LANG)), superscript=True)
        if gr_cnt > 1:
            _generate_text(par, subscript(gr_cnt, TO_LANG), subscript=True)


def _generate_usage(par, u: Alignment) -> None:
    _generate_index(par, u)
    if not u.has_alternatives():
        return

    _generate_text(par, f" {CF_SEP}")
    if u.orig.main_alt:
        source = main_source(u.orig.lang, u.idx.data[2] == "W")
        _generate_usage_alt_main(par, u.orig, source)

    if u.orig.var_alt:
        _generate_usage_alt_vars(par, u.orig)

    if u.trans.main_alt:
        source = main_source(u.trans.lang, u.idx.data[2] == "W")
        _generate_usage_alt_main(par, u.trans, source)

    if u.trans.var_alt:
        _generate_usage_alt_vars(par, u.trans)


def docx_result(par, usage: SortedSet, src_style: str) -> None:
    """
    usage: list of indices of usages also containing their styles
    """
    first = True
    for next in usage:
        if first:
            first = False
        else:
            _generate_text(par, "; ")
        try:
            _generate_usage(par, next)
        except Exception as e:
            log.error(
                f"При генериране възникна проблем в ред {next.idx} или групата му"
            )
            log.error(e)
            continue


def _generate_counts(par, d: Union[SortedDict, dict], trans: bool = False) -> None:
    try:
        c = Counter.get_dict_counts(d).get_counts(trans)
    except StopIteration:
        keys = []
        key = next(iter(d))
        if key:
            keys += [key]
        while key and key in d and bool(d[key]):
            d = d[key]
            key = next(iter(d))
            if key:
                keys += [key]
        # TODO: Better reporting
        if keys:
            log.error(f"При генериране възникна проблем с {'|'.join(keys)}")
        else:
            log.error(f"При генериране възникна неидентифициран проблем")
        log.error(f"При генериране неуспешно преброяване в {d}")
        return
    run = par.add_run()
    run.add_text(" (")
    if c[0]:
        run.add_text(str(c[0]))
        if c[1]:
            run.add_text(" + ")
    if c[1]:
        run.add_text(str(c[1]))
        _generate_text(par, "var", superscript=True)


def _generate_usage_line(lang: str, d: SortedDict, doc: Document) -> None:
    """Merges together all occurences for the purposes of ordering, because usage pairs are not shown"""
    trans_lang = TO_LANG if lang == FROM_LANG else FROM_LANG
    for t, bottom_d in d.items():
        # c = _get_dict_counts(d).get_counts(True)
        # if not c[0] and not c[1]:
        #    return

        par = doc.add_paragraph()
        par.style = doc.styles[BULLET_STYLE]
        par.style.font.name = GENERIC_FONT
        par.paragraph_format.line_spacing = Pt(12)
        par.paragraph_format.space_before = Cm(0)
        par.paragraph_format.space_after = Cm(0)
        par.paragraph_format.left_indent = Cm(LEVEL_OFFSET * 4)
        # ft = t if t[0] in SPECIAL_CHARS else f"{BULLET_CH} {t}"

        _generate_text(par, t, fonts[trans_lang])
        _generate_counts(par, bottom_d, True)

        run = par.add_run()
        # run.font.name = GENERIC_FONT
        run.add_text("): ")
        total = SortedSet()
        for nxt in bottom_d.values():
            total.update(nxt)
        docx_result(par, total, lang)


def _generate_line(level: int, lang: str, d: SortedDict, doc: Document) -> None:
    """Builds the hierarchical entries of the dictionary.
    Recursion ensures that this works with variable depth.

    Args:
        level (int): 0-indexed depth, runs along with dict depth
        lang (str): original language
        d (SortedDict): level of dictionary to be exported
        doc (Document): export document
    """
    for li, next_d in d.items():
        if level == 0 and not li:
            continue
        # c = _get_dict_counts(d).get_counts(False)
        # if not c[0] and not c[1]:
        #    return
        if li:
            par = doc.add_paragraph()
            par.style.font.name = GENERIC_FONT
            # TODO: Less distance between lines, make exact number
            par.paragraph_format.space_before = Cm(0)
            par.paragraph_format.space_after = Cm(0)
            if level > 0:
                par.paragraph_format.first_line_indent = Cm(LEVEL_OFFSET * level)

            prefix = "" if li[0] in SPECIAL_CHARS else f"{INDENT_CH} " * level
            _generate_text(
                par,
                f"{prefix}{li}",
                fonts[lang],
                size=Pt(12 if level == 0 else 11),
                bold=level == 0,
                indent=Cm(LEVEL_OFFSET * level),
            )

            _generate_counts(par, next_d, False)
            _generate_text(par, ")")

        try:
            any_of_any = any_grandchild(next_d)
        except AssertionError as ae:
            log.error(
                f"При генериране възникна проблем с една от {len(next_d)} употреби на {li}"
            )
            raise ae
        if type(any_of_any) is SortedSet:  # bottom of structure
            _generate_usage_line(lang, next_d, doc)
        else:
            try:
                _generate_line(level + 1, lang, next_d, doc)
            except AssertionError:
                log.error(
                    f"При генериране възникна проблем с една от {len(next_d)} употреби на {li}"
                )


def generate_docx(d: SortedDict, lang: str, fname: str) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = GENERIC_FONT
    _generate_line(0, lang, d, doc)

    doc.save(fname)
